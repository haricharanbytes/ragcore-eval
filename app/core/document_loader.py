import logging
from pathlib import Path

from langchain_core.documents import Document as LCDocument
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentLoadError(Exception):
    """Raised when a file can't be parsed"""
    pass


def _load_pdf(file_path: Path) -> list[LCDocument]:
    documents = []
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:
        raise DocumentLoadError(f"Could not open PDF (corrupted or encrypted?): {exc}") from exc

    if len(reader.pages) == 0:
        raise DocumentLoadError("PDF has no pages.")

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                LCDocument(
                    page_content=text,
                    metadata={"source": file_path.name, "page": page_num},
                )
            )

    if not documents:
        raise DocumentLoadError(
            "No extractable text found — this may be a scanned/image-only PDF."
        )
    return documents


def _load_docx(file_path: Path) -> list[LCDocument]:
    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        raise DocumentLoadError(f"Could not open DOCX file: {exc}") from exc

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        raise DocumentLoadError("DOCX file appears to be empty.")

    # DOCX has no reliable page concept, so we treat the whole file as
    # a single logical unit — chunking splits it further downstream.
    return [LCDocument(page_content=full_text, metadata={"source": file_path.name, "page": None})]


def _load_txt(file_path: Path) -> list[LCDocument]:
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        raise DocumentLoadError(f"Could not read text file: {exc}") from exc

    if not text.strip():
        raise DocumentLoadError("Text file is empty.")

    return [LCDocument(page_content=text, metadata={"source": file_path.name, "page": None})]


# Dispatch table: extension -> loader function.
# Keeping this as data (not if/elif) makes it trivial to extend.
_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_txt,
}


def load_document(file_path: Path) -> list[LCDocument]:
    """
    Load a single file into LangChain Documents.

    Raises:
        DocumentLoadError: unsupported extension, corrupt file, or no
            extractable text — always a clear, user-facing message.
    """
    extension = file_path.suffix.lower()
    loader_fn = _LOADERS.get(extension)

    if loader_fn is None:
        raise DocumentLoadError(
            f"Unsupported file type '{extension}'. Supported: {', '.join(_LOADERS)}"
        )

    logger.info("Loading document: %s", file_path.name)
    documents = loader_fn(file_path)
    logger.info("Loaded %d page/section(s) from %s", len(documents), file_path.name)
    return documents